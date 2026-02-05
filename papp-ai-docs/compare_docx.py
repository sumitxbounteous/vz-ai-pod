import base64
import difflib
import io
from docx import Document
import diff_match_patch as dmp_module
import mammoth

def compare_io_files(doc1_base64, doc2_base64):
    # Decode base64 strings to bytes
    doc1_bytes = base64.b64decode(doc1_base64)
    doc2_bytes = base64.b64decode(doc2_base64)

    # Load both documents from bytes
    doc1 = Document(io.BytesIO(doc1_bytes))
    doc2 = Document(io.BytesIO(doc2_bytes))

    # Extract text from paragraphs
    doc1_text = [para.text for para in doc1.paragraphs]
    doc2_text = [para.text for para in doc2.paragraphs]
    dmp = dmp_module.diff_match_patch()
    ret_diff_html: list[str] = []
    ret_summary: list[str] = []
    for para1, para2 in zip(doc1_text, doc2_text):
        diffs = dmp.diff_main(para1, para2)
        dmp.diff_cleanupEfficiency(diffs)
        diff_html = dmp.diff_prettyHtml(diffs)
        ret_diff_html.append(diff_html)
        for op, data in diffs:
            if op == dmp_module.diff_match_patch.DIFF_INSERT:
                ret_summary.append(f"Added: {data}")    
            elif op == dmp_module.diff_match_patch.DIFF_DELETE:
                ret_summary.append(f"Deleted: {data}")

    doc1_html = mammoth.convert_to_html(io.BytesIO(doc1_bytes))
    doc2_html = mammoth.convert_to_html(io.BytesIO(doc2_bytes))
      # The generated HTML
    # messages = result.messages  # Any warnings or messages
    return ret_diff_html, ret_summary, doc1_html.value, doc2_html.value


# def compare_word_documents(doc1_path, doc2_path):
#     # Load both documents
#     doc1 = Document(doc1_path)
#     doc2 = Document(doc2_path)

#     # Extract text from paragraphs
#     doc1_text = [para.text for para in doc1.paragraphs]
#     doc2_text = [para.text for para in doc2.paragraphs]

#     # Use difflib for comparison
#     differ = difflib.Differ()
#     diff = list(differ.compare(doc1_text, doc2_text))

#     # Count changes
#     # additions = sum(1 for line in diff if line.startswith('+ '))
#     # deletions = sum(1 for line in diff if line.startswith('- '))

#     # print(f"Total additions: {additions}")
#     # print(f"Total deletions: {deletions}")

#     # Display differences
#     # for line in diff:
#     #     if line.startswith('+ ') or line.startswith('- '):
#     #         print(line)

#     return (diff, doc1.paragraphs, doc2.paragraphs)
